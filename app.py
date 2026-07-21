# -*- coding: utf-8 -*-
"""
电化学实验条件数据库平台 - Electrochemical Experimental Conditions Database Platform
"""
import os
import json
import csv
import io
import math
import re
from datetime import datetime, timezone
from io import StringIO, BytesIO

from flask import (
    Flask, render_template, redirect, url_for, flash, request,
    jsonify, send_file, session, make_response
)
from flask_login import (
    LoginManager, login_user, logout_user, login_required, current_user
)
from sqlalchemy import or_, func, case
from werkzeug.utils import secure_filename

from models import db, User, Experiment, Electrode, Electrolyte, Sample, Instrument, Analysis, AuditLog, Literature, literature_experiment, ReactionFingerprint
from models import experiment_electrode, experiment_electrolyte, experiment_sample
from forms import (
    LoginForm, RegisterForm, ExperimentForm, AnalysisForm, SearchForm, ImportForm, LiteratureForm, ReactionSearchForm
)

# ------- App Factory -------




app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'ecdb-secret-key-change-in-production')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///ecdb.sqlite3')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB max upload

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db.init_app(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = '请先登录以访问此页面。'

@app.context_processor
def inject_globals():
    """Make common models and utilities available in all templates."""
    return dict(
        Literature=Literature,
        now=datetime.now,
    )



@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))


# ------- Helpers -------

def log_audit(action, table_name, record_id=None, details=None):
    """Create an audit log entry."""
    if current_user and current_user.is_authenticated:
        log = AuditLog(
            user_id=current_user.id,
            action=action,
            table_name=table_name,
            record_id=record_id,
            details=details,
            ip_address=request.remote_addr
        )
        db.session.add(log)
        db.session.commit()


def get_or_create_instrument(model_name, manufacturer=None):
    """Get or create an instrument record."""
    inst = Instrument.query.filter_by(model=model_name).first()
    if not inst:
        inst = Instrument(model=model_name, manufacturer=manufacturer or '')
        db.session.add(inst)
        db.session.flush()
    return inst.id


def get_or_create_electrode(mat, etype, area=None, geometry=None, pretreatment=None):
    """Get or create an electrode record."""
    if not mat:
        return None
    elec = Electrode.query.filter_by(material=mat, electrode_type=etype).first()
    if not elec:
        elec = Electrode(
            electrode_type=etype, material=mat,
            area=area, geometry=geometry, pretreatment=pretreatment
        )
        db.session.add(elec)
        db.session.flush()
    else:
        # Update key fields
        if area is not None:
            elec.area = area
        if geometry:
            elec.geometry = geometry
        if pretreatment:
            elec.pretreatment = pretreatment
    return elec.id


def get_or_create_electrolyte(solvent, salt=None, conc=None, ph=None):
    """Get or create an electrolyte record."""
    if not solvent:
        return None
    elec = Electrolyte.query.filter_by(solvent=solvent, supporting_salt=salt or '').first()
    if not elec:
        elec = Electrolyte(
            solvent=solvent, supporting_salt=salt or '',
            concentration=conc, ph=ph
        )
        db.session.add(elec)
        db.session.flush()
    return elec.id


def get_or_create_sample(name, conc=None, role='substrate'):
    """Get or create a sample record."""
    if not name:
        return None
    samp = Sample.query.filter_by(name=name).first()
    if not samp:
        samp = Sample(name=name, concentration=conc)
        db.session.add(samp)
        db.session.flush()
    return samp.id


def calculate_faraday_efficiency(yield_pct, total_charge, electrons_per_mol=2):
    """
    Calculate Faraday efficiency.
    FE(%) = (n * F * product_mol) / total_charge * 100
    Where F = 96485 C/mol
    """
    if not total_charge or total_charge <= 0:
        return None
    # If yield_pct is already a percentage of theoretical max,
    # we need actual product moles. We'll use a simplified version.
    return yield_pct  # Placeholder - real calculation needs product quantification


def calculate_rsd(mean, std):
    """Calculate relative standard deviation."""
    if not mean or mean == 0:
        return None
    return (std / mean) * 100


# ------- Routes: Auth -------

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and user.check_password(form.password.data):
            login_user(user)
            log_audit('LOGIN', 'user', user.id, f'User {user.username} logged in')
            flash('登录成功！', 'success')
            next_page = request.args.get('next')
            return redirect(next_page or url_for('index'))
        flash('用户名或密码错误', 'danger')
    return render_template('login.html', form=form)


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = RegisterForm()
    if form.validate_on_submit():
        if User.query.filter_by(username=form.username.data).first():
            flash('用户名已存在', 'danger')
            return render_template('register.html', form=form)
        user = User(
            username=form.username.data,
            email=form.email.data,
            affiliation=form.affiliation.data
        )
        user.set_password(form.password.data)
        db.session.add(user)
        db.session.commit()
        flash('注册成功！请登录。', 'success')
        return redirect(url_for('login'))
    return render_template('register.html', form=form)


@app.route('/logout')
@login_required
def logout():
    log_audit('LOGOUT', 'user', current_user.id, f'User {current_user.username} logged out')
    logout_user()
    flash('已退出登录。', 'info')
    return redirect(url_for('login'))


# ------- Routes: Dashboard -------

@app.route('/')
@login_required
def index():
    total_experiments = Experiment.query.filter_by(is_active=True).count()
    total_literature = Literature.query.count()
    total_literature = Literature.query.count()
    total_electrodes = Electrode.query.count()
    total_electrolytes = Electrolyte.query.count()
    
    # Recent experiments
    recent = Experiment.query.filter_by(is_active=True)\
        .order_by(Experiment.created_at.desc()).limit(5).all()
    
    # Mode distribution
    mode_counts = db.session.query(
        Experiment.electro_mode, func.count(Experiment.id)
    ).filter(Experiment.is_active == True)\
     .group_by(Experiment.electro_mode).all()
    
    # Material distribution (top 10 WE materials)
    we_materials = db.session.query(
        Electrode.material, func.count(experiment_electrode.c.experiment_id).label('cnt')
    ).join(experiment_electrode, Electrode.id == experiment_electrode.c.electrode_id)\
     .filter(Electrode.electrode_type == 'WE')\
     .group_by(Electrode.material)\
     .order_by(func.count(experiment_electrode.c.experiment_id).desc())\
     .limit(10).all()
    
    return render_template('index.html',
        total_experiments=total_experiments,
        total_electrodes=total_electrodes,
        total_electrolytes=total_electrolytes,
        total_literature=total_literature,
        recent=recent,
        mode_counts=mode_counts,
        we_materials=we_materials
    )


# ------- Routes: Experiments -------

@app.route('/experiments')
@login_required
def experiment_list():
    page = request.args.get('page', 1, type=int)
    per_page = 20
    form = SearchForm()
    
    query = Experiment.query.filter_by(is_active=True)
    
    # Apply filters
    we_mat = request.args.get('we_material', '')
    solvent = request.args.get('solvent', '')
    substrate = request.args.get('substrate', '')
    electrolyte = request.args.get('electrolyte', '')
    electro_mode = request.args.get('electro_mode', '')
    ref_type = request.args.get('reference_type', '')
    sim_smiles = request.args.get('sim_smiles', '')
    sim_threshold = request.args.get('sim_threshold', 0.0, type=float)
    ph_min = request.args.get('ph_min', '', type=float)
    ph_max = request.args.get('ph_max', '', type=float)
    temp_min = request.args.get('temp_min', '', type=float)
    temp_max = request.args.get('temp_max', '', type=float)
    keyword = request.args.get('keyword', '')
    
    if we_mat:
        query = query.join(experiment_electrode).join(Electrode)\
            .filter(Electrode.electrode_type == 'WE',
                    Electrode.material.ilike(f'%{we_mat}%'))
    if solvent:
        query = query.join(experiment_electrolyte).join(Electrolyte)\
            .filter(Electrolyte.solvent.ilike(f'%{solvent}%'))
    if substrate:
        query = query.join(experiment_sample).join(Sample)\
            .filter(Sample.name.ilike(f'%{substrate}%'))
    if electrolyte:
        query = query.join(experiment_electrolyte).join(Electrolyte)\
            .filter(Electrolyte.supporting_salt.ilike(f'%{electrolyte}%'))
    if electro_mode:
        query = query.filter(Experiment.electro_mode == electro_mode)
    if ref_type:
        query = query.filter(Experiment.reference_type == ref_type)
    if ph_min is not None:
        query = query.filter(Experiment.ph >= ph_min)
    if ph_max is not None:
        query = query.filter(Experiment.ph <= ph_max)
    if temp_min is not None:
        query = query.filter(Experiment.temperature >= temp_min)
    if temp_max is not None:
        query = query.filter(Experiment.temperature <= temp_max)
    if keyword:
        query = query.filter(or_(
            Experiment.name.ilike(f'%{keyword}%'),
            Experiment.notes.ilike(f'%{keyword}%'),
            Experiment.atmosphere.ilike(f'%{keyword}%'),
            Experiment.smiles.ilike(f'%{keyword}%')
        ))
    
    experiments = query.order_by(Experiment.created_at.desc())\
        .paginate(page=page, per_page=per_page, error_out=False)
    
    return render_template('experiments/list.html',
        experiments=experiments, form=form, 
        request_args=request.args
    )


@app.route('/experiments/create', methods=['GET', 'POST'])
@login_required
def experiment_create():
    form = ExperimentForm()
    if form.validate_on_submit():
        try:
            # Create/Get instrument
            instrument_id = None
            if form.instrument_model.data:
                instrument_id = get_or_create_instrument(form.instrument_model.data)
            
            # Create experiment record
            exp = Experiment(
                name=form.name.data,
                date_time=form.date_time.data or datetime.now(timezone.utc),
                user_id=current_user.id,
                batch_id=form.batch_id.data,
                instrument_id=instrument_id,
                # Reactor
                cell_type=form.cell_type.data,
                cell_volume=form.cell_volume.data,
                separator=form.separator.data,
                # Solution
                ph=form.ph.data,
                temperature=form.temperature.data,
                stirring_rpm=form.stirring_rpm.data,
                atmosphere=form.atmosphere.data,
                feed_mode=form.feed_mode.data,
                feed_rate=form.feed_rate.data,
                # Electrochemical
                electro_mode=form.electro_mode.data,
                current=form.current.data,
                potential=form.potential.data,
                scan_rate=form.scan_rate.data,
                total_charge=form.total_charge.data,
                duration=form.duration.data,
                duration_unit=form.duration_unit.data,
                # Light
                light_source=form.light_source.data,
                light_intensity=form.light_intensity.data,
                light_duration=form.light_duration.data,
                # Measurement
                measurement_method=form.measurement_method.data,
                sampling_rate=form.sampling_rate.data,
                analysis_method=form.analysis_method.data,
                uncertainty=form.uncertainty.data,
                smiles=form.smiles.data,
                substrate_smiles=form.substrate_smiles.data,
                product_smiles=form.product_smiles.data,
                reference_type=form.reference_type.data,
                doi=form.doi.data,
                patent_number=form.patent_number.data,
                # Other
                equipment_model=form.instrument_model.data,
                operator_name=form.operator_name.data or current_user.username,
                safety_info=form.safety_info.data,
                waste_category=form.waste_category.data,
                notes=form.notes.data
            )
            db.session.add(exp)
            db.session.flush()
            
            # Create and associate electrodes
            if form.we_material.data:
                eid = get_or_create_electrode(
                    form.we_material.data, 'WE',
                    area=form.we_area.data,
                    geometry=form.we_geometry.data,
                    pretreatment=form.we_pretreatment.data
                )
                if eid:
                    stmt = experiment_electrode.insert().values(
                        experiment_id=exp.id, electrode_id=eid, function='WE'
                    )
                    db.session.execute(stmt)
            
            if form.ce_material.data:
                eid = get_or_create_electrode(form.ce_material.data, 'CE')
                if eid:
                    stmt = experiment_electrode.insert().values(
                        experiment_id=exp.id, electrode_id=eid, function='CE'
                    )
                    db.session.execute(stmt)
            
            if form.re_type.data:
                eid = get_or_create_electrode(form.re_type.data, 'RE')
                if eid:
                    stmt = experiment_electrode.insert().values(
                        experiment_id=exp.id, electrode_id=eid, function='RE'
                    )
                    db.session.execute(stmt)
            
            # Create and associate electrolyte
            if form.solvent.data:
                eid = get_or_create_electrolyte(
                    form.solvent.data, form.electrolyte.data,
                    conc=form.electrolyte_conc.data, ph=form.ph.data
                )
                if eid:
                    stmt = experiment_electrolyte.insert().values(
                        experiment_id=exp.id, electrolyte_id=eid
                    )
                    db.session.execute(stmt)
            
            # Create and associate samples
            if form.substrate.data:
                sid = get_or_create_sample(form.substrate.data, form.substrate_conc.data, 'substrate')
                if sid:
                    stmt = experiment_sample.insert().values(
                        experiment_id=exp.id, sample_id=sid, role='substrate'
                    )
                    db.session.execute(stmt)
            
            if form.catalyst.data:
                sid = get_or_create_sample(form.catalyst.data, form.catalyst_load.data, 'catalyst')
                if sid:
                    stmt = experiment_sample.insert().values(
                        experiment_id=exp.id, sample_id=sid, role='catalyst'
                    )
                    db.session.execute(stmt)
            
            if form.additives.data:
                sid = get_or_create_sample(form.additives.data, role='additive')
                if sid:
                    stmt = experiment_sample.insert().values(
                        experiment_id=exp.id, sample_id=sid, role='additive'
                    )
                    db.session.execute(stmt)
            
            db.session.commit()
            log_audit('CREATE', 'experiment', exp.id, f'Created experiment: {exp.name}')
            flash(f'实验 "{exp.name}" 创建成功！', 'success')
            return redirect(url_for('experiment_detail', experiment_id=exp.id))
            
        except Exception as e:
            db.session.rollback()
            flash(f'创建实验失败: {str(e)}', 'danger')
    
    return render_template('experiments/create.html', form=form)


@app.route('/experiments/<int:experiment_id>')
@login_required
def experiment_detail(experiment_id):
    exp = db.session.get(Experiment, experiment_id)
    if not exp or not exp.is_active:
        flash('实验未找到。', 'warning')
        return redirect(url_for('experiment_list'))
    
    # Get electrodes by function
    we_electrode = None
    ce_electrode = None
    re_electrode = None
    for elec in exp.electrodes:
        funcs = db.session.execute(
            db.select(experiment_electrode.c.function).where(
                experiment_electrode.c.experiment_id == exp.id,
                experiment_electrode.c.electrode_id == elec.id
            )
        ).scalars().all()
        for f in funcs:
            if f == 'WE':
                we_electrode = elec
            elif f == 'CE':
                ce_electrode = elec
            elif f == 'RE':
                re_electrode = elec
    
    analyses = exp.analyses.order_by(Analysis.created_at.desc()).all()
    
    return render_template('experiments/detail.html',
        experiment=exp, we_electrode=we_electrode,
        ce_electrode=ce_electrode, re_electrode=re_electrode,
        analyses=analyses
    )


@app.route('/experiments/<int:experiment_id>/edit', methods=['GET', 'POST'])
@login_required
def experiment_edit(experiment_id):
    exp = db.session.get(Experiment, experiment_id)
    if not exp or not exp.is_active:
        flash('实验未找到。', 'warning')
        return redirect(url_for('experiment_list'))
    
    form = ExperimentForm(obj=exp)
    
    # Pre-populate electrode/electrolyte/sample fields
    if request.method == 'GET':
        for elec in exp.electrodes:
            funcs = db.session.execute(
                db.select(experiment_electrode.c.function).where(
                    experiment_electrode.c.experiment_id == exp.id,
                    experiment_electrode.c.electrode_id == elec.id
                )
            ).scalars().all()
            for f in funcs:
                if f == 'WE':
                    form.we_material.data = elec.material
                    form.we_area.data = elec.area
                    form.we_geometry.data = elec.geometry
                    form.we_pretreatment.data = elec.pretreatment
                elif f == 'CE':
                    form.ce_material.data = elec.material
                elif f == 'RE':
                    form.re_type.data = elec.material
        for elec in exp.electrolytes:
            form.solvent.data = elec.solvent
            form.electrolyte.data = elec.supporting_salt
            form.electrolyte_conc.data = elec.concentration
        for samp in exp.samples:
            # Check role
            roles = db.session.execute(
                db.select(experiment_sample.c.role).where(
                    experiment_sample.c.experiment_id == exp.id,
                    experiment_sample.c.sample_id == samp.id
                )
            ).scalars().all()
            for r in roles:
                if r == 'substrate':
                    form.substrate.data = samp.name
                    form.substrate_conc.data = samp.concentration
                elif r == 'catalyst':
                    form.catalyst.data = samp.name
                    form.catalyst_load.data = samp.concentration
                elif r == 'additive':
                    form.additives.data = samp.name
    
    if form.validate_on_submit():
        try:
            exp.name = form.name.data
            exp.date_time = form.date_time.data
            exp.batch_id = form.batch_id.data
            exp.cell_type = form.cell_type.data
            exp.cell_volume = form.cell_volume.data
            exp.separator = form.separator.data
            exp.ph = form.ph.data
            exp.temperature = form.temperature.data
            exp.stirring_rpm = form.stirring_rpm.data
            exp.atmosphere = form.atmosphere.data
            exp.feed_mode = form.feed_mode.data
            exp.feed_rate = form.feed_rate.data
            exp.electro_mode = form.electro_mode.data
            exp.current = form.current.data
            exp.potential = form.potential.data
            exp.scan_rate = form.scan_rate.data
            exp.total_charge = form.total_charge.data
            exp.duration = form.duration.data
            exp.duration_unit = form.duration_unit.data
            exp.light_source = form.light_source.data
            exp.light_intensity = form.light_intensity.data
            exp.light_duration = form.light_duration.data
            exp.measurement_method = form.measurement_method.data
            exp.sampling_rate = form.sampling_rate.data
            exp.analysis_method = form.analysis_method.data
            exp.uncertainty = form.uncertainty.data
            exp.equipment_model = form.instrument_model.data
            exp.operator_name = form.operator_name.data
            exp.safety_info = form.safety_info.data
            exp.waste_category = form.waste_category.data
            exp.notes = form.notes.data
            exp.smiles = form.smiles.data
            exp.substrate_smiles = form.substrate_smiles.data
            exp.product_smiles = form.product_smiles.data
            exp.reference_type = form.reference_type.data
            exp.doi = form.doi.data
            exp.patent_number = form.patent_number.data
            exp.version += 1
            
            db.session.commit()
            log_audit('UPDATE', 'experiment', exp.id, f'Updated experiment: {exp.name}')
            flash('实验更新成功！', 'success')
            return redirect(url_for('experiment_detail', experiment_id=exp.id))
        except Exception as e:
            db.session.rollback()
            flash(f'更新失败: {str(e)}', 'danger')
    
    return render_template('experiments/create.html', form=form, edit=True, experiment=exp)


@app.route('/experiments/<int:experiment_id>/delete', methods=['POST'])
@login_required
def experiment_delete(experiment_id):
    exp = db.session.get(Experiment, experiment_id)
    if not exp:
        flash('实验未找到。', 'warning')
        return redirect(url_for('experiment_list'))
    
    exp.is_active = False
    db.session.commit()
    log_audit('DELETE', 'experiment', exp.id, f'Soft deleted experiment: {exp.name}')
    flash('实验已删除。', 'info')
    return redirect(url_for('experiment_list'))


# ------- Routes: Analysis -------

@app.route('/experiments/<int:experiment_id>/analysis/create', methods=['GET', 'POST'])
@login_required
def analysis_create(experiment_id):
    exp = db.session.get(Experiment, experiment_id)
    if not exp or not exp.is_active:
        flash('实验未找到。', 'warning')
        return redirect(url_for('experiment_list'))
    
    form = AnalysisForm()
    if form.validate_on_submit():
        try:
            analysis = Analysis(
                experiment_id=exp.id,
                yield_pct=form.yield_pct.data,
                selectivity=form.selectivity.data,
                faraday_efficiency=form.faraday_efficiency.data,
                byproducts=form.byproducts.data,
                reproducibility_mean=form.reproducibility_mean.data,
                reproducibility_std=form.reproducibility_std.data,
                num_replicates=form.num_replicates.data or 1,
                anomalies=form.anomalies.data,
                notes=form.notes.data,
                raw_data_path=form.raw_data_path.data
            )
            
            # Auto-calculate RSD
            if analysis.reproducibility_mean and analysis.reproducibility_std:
                analysis.reproducibility_rsd = calculate_rsd(
                    analysis.reproducibility_mean, analysis.reproducibility_std
                )
            
            db.session.add(analysis)
            db.session.commit()
            log_audit('CREATE', 'analysis', analysis.id, f'Added analysis for experiment #{exp.id}')
            flash('分析结果已保存！', 'success')
            return redirect(url_for('experiment_detail', experiment_id=exp.id))
        except Exception as e:
            db.session.rollback()
            flash(f'保存失败: {str(e)}', 'danger')
    
    return render_template('analysis/create.html', form=form, experiment=exp)


@app.route('/analysis/<int:analysis_id>/edit', methods=['GET', 'POST'])
@login_required
def analysis_edit(analysis_id):
    analysis = db.session.get(Analysis, analysis_id)
    if not analysis:
        flash('分析记录未找到。', 'warning')
        return redirect(url_for('index'))
    
    form = AnalysisForm(obj=analysis)
    if form.validate_on_submit():
        try:
            analysis.yield_pct = form.yield_pct.data
            analysis.selectivity = form.selectivity.data
            analysis.faraday_efficiency = form.faraday_efficiency.data
            analysis.byproducts = form.byproducts.data
            analysis.reproducibility_mean = form.reproducibility_mean.data
            analysis.reproducibility_std = form.reproducibility_std.data
            analysis.num_replicates = form.num_replicates.data or 1
            analysis.anomalies = form.anomalies.data
            analysis.notes = form.notes.data
            analysis.raw_data_path = form.raw_data_path.data
            
            if analysis.reproducibility_mean and analysis.reproducibility_std:
                analysis.reproducibility_rsd = calculate_rsd(
                    analysis.reproducibility_mean, analysis.reproducibility_std
                )
            
            db.session.commit()
            log_audit('UPDATE', 'analysis', analysis.id, 'Updated analysis')
            flash('分析结果已更新！', 'success')
            return redirect(url_for('experiment_detail', experiment_id=analysis.experiment_id))
        except Exception as e:
            db.session.rollback()
            flash(f'更新失败: {str(e)}', 'danger')
    
    return render_template('analysis/create.html', form=form, experiment=analysis.experiment, analysis=analysis)


# ------- Routes: Statistics & Visualization -------

@app.route('/statistics')
@login_required
def statistics():
    # Summary stats
    total_exp = Experiment.query.filter_by(is_active=True).count()
    with_analysis = Analysis.query.count()
    
    # Average yield by electrode material
    yield_by_we = db.session.query(
        Electrode.material,
        func.avg(Analysis.yield_pct).label('avg_yield'),
        func.count(Analysis.id).label('count')
    ).select_from(Electrode)\
     .join(experiment_electrode, Electrode.id == experiment_electrode.c.electrode_id)\
     .join(Experiment, Experiment.id == experiment_electrode.c.experiment_id)\
     .join(Analysis, Analysis.experiment_id == Experiment.id)\
     .filter(Electrode.electrode_type == 'WE',
             Experiment.is_active == True,
             Analysis.yield_pct != None)\
     .group_by(Electrode.material)\
     .order_by(func.avg(Analysis.yield_pct).desc()).all()
    
    # Average FE by electrolyte
    fe_by_electrolyte = db.session.query(
        Electrolyte.supporting_salt,
        func.avg(Analysis.faraday_efficiency).label('avg_fe'),
        func.count(Analysis.id).label('count')
    ).select_from(Electrolyte)\
     .join(experiment_electrolyte, Electrolyte.id == experiment_electrolyte.c.electrolyte_id)\
     .join(Experiment, Experiment.id == experiment_electrolyte.c.experiment_id)\
     .join(Analysis, Analysis.experiment_id == Experiment.id)\
     .filter(Experiment.is_active == True,
             Analysis.faraday_efficiency != None)\
     .group_by(Electrolyte.supporting_salt)\
     .order_by(func.avg(Analysis.faraday_efficiency).desc()).all()
    
    # Experiments over time
    exp_over_time = db.session.query(
        func.date(Experiment.date_time).label('date'),
        func.count(Experiment.id).label('count')
    ).filter(Experiment.is_active == True)\
     .group_by(func.date(Experiment.date_time))\
     .order_by(func.date(Experiment.date_time)).all()
    
    # Mode distribution for chart
    mode_data = db.session.query(
        Experiment.electro_mode,
        func.count(Experiment.id).label('count')
    ).filter(Experiment.is_active == True)\
     .group_by(Experiment.electro_mode).all()
    
    # Temperature vs Yield scatter (for chart)
    temp_yield_data = db.session.query(
        Experiment.temperature,
        Analysis.yield_pct
    ).join(Analysis, Analysis.experiment_id == Experiment.id)\
     .filter(Experiment.is_active == True,
             Experiment.temperature != None,
             Analysis.yield_pct != None).all()
    
    # pH vs FE scatter
    ph_fe_data = db.session.query(
        Experiment.ph,
        Analysis.faraday_efficiency
    ).join(Analysis, Analysis.experiment_id == Experiment.id)\
     .filter(Experiment.is_active == True,
             Experiment.ph != None,
             Analysis.faraday_efficiency != None).all()
    
    # Convert for JSON serialization
    chart_data = {
        'mode_labels': [r[0] or '未指定' for r in mode_data],
        'mode_values': [int(r[1]) for r in mode_data],
        'temp_yield': [[float(r[0]), float(r[1])] for r in temp_yield_data if r[0] and r[1]],
        'ph_fe': [[float(r[0]), float(r[1])] for r in ph_fe_data if r[0] and r[1]],
        'we_yield_labels': [r[0] for r in yield_by_we],
        'we_yield_values': [round(float(r[1]), 1) for r in yield_by_we],
        'fe_elec_labels': [r[0] or '未指定' for r in fe_by_electrolyte],
        'fe_elec_values': [round(float(r[1]), 1) for r in fe_by_electrolyte],
        'dates': [str(r[0]) for r in exp_over_time],
        'exp_counts': [int(r[1]) for r in exp_over_time]
    }
    
    return render_template('analysis/statistics.html',
        total_exp=total_exp,
        with_analysis=with_analysis,
        yield_by_we=yield_by_we,
        fe_by_electrolyte=fe_by_electrolyte,
        chart_data=json.dumps(chart_data)
    )


# =====================================================================
# Routes: Literature / Patents  (文献/专利管理)
# =====================================================================

@app.route('/literature')
@login_required
def literature_list():
    page = request.args.get('page', 1, type=int)
    per_page = 20
    lit_type = request.args.get('type', '')
    keyword = request.args.get('keyword', '')
    year_from = request.args.get('year_from', type=int)
    query = Literature.query
    if lit_type:
        query = query.filter(Literature.literature_type == lit_type)
    if keyword:
        query = query.filter(db.or_(
            Literature.title.ilike(f'%{keyword}%'),
            Literature.authors.ilike(f'%{keyword}%'),
            Literature.keywords.ilike(f'%{keyword}%'),
            Literature.abstract.ilike(f'%{keyword}%')
        ))
    if year_from:
        query = query.filter(Literature.year >= year_from)
    literatures = query.order_by(Literature.created_at.desc())\
        .paginate(page=page, per_page=per_page, error_out=False)
    return render_template('literature/list.html', literatures=literatures)


@app.route('/literature/create', methods=['GET', 'POST'])
@login_required
def literature_create():
    form = LiteratureForm()
    if form.validate_on_submit():
        try:
            lit = Literature(
                title=form.title.data,
                authors=form.authors.data,
                journal=form.journal.data,
                year=form.year.data,
                volume=form.volume.data,
                issue=form.issue.data,
                pages=form.pages.data,
                doi=form.doi.data,
                patent_number=form.patent_number.data,
                literature_type=form.literature_type.data,
                abstract=form.abstract.data,
                keywords=form.keywords.data,
                url=form.url.data,
                research_field=form.research_field.data,
                reaction_type=form.reaction_type.data,
                notes=form.notes.data,
                user_id=current_user.id
            )
            db.session.add(lit)
            db.session.commit()
            log_audit('CREATE', 'literature', lit.id, f'Added literature: {lit.title}')
            flash('文献添加成功！', 'success')
            return redirect(url_for('literature_list'))
        except Exception as e:
            db.session.rollback()
            flash(f'添加失败: {str(e)}', 'danger')
    return render_template('literature/create.html', form=form)


@app.route('/literature/<int:lit_id>')
@login_required
def literature_detail(lit_id):
    lit = db.session.get(Literature, lit_id)
    if not lit:
        flash('文献未找到。', 'warning')
        return redirect(url_for('literature_list'))
    return render_template('literature/detail.html', lit=lit)


@app.route('/literature/<int:lit_id>/edit', methods=['GET', 'POST'])
@login_required
def literature_edit(lit_id):
    lit = db.session.get(Literature, lit_id)
    if not lit:
        flash('文献未找到。', 'warning')
        return redirect(url_for('literature_list'))
    form = LiteratureForm(obj=lit)
    if form.validate_on_submit():
        try:
            form.populate_obj(lit)
            db.session.commit()
            log_audit('UPDATE', 'literature', lit.id, f'Updated: {lit.title}')
            flash('文献信息已更新。', 'success')
            return redirect(url_for('literature_detail', lit_id=lit.id))
        except Exception as e:
            db.session.rollback()
            flash(f'更新失败: {str(e)}', 'danger')
    return render_template('literature/create.html', form=form, edit=True, lit=lit)


@app.route('/literature/<int:lit_id>/delete', methods=['POST'])
@login_required
def literature_delete(lit_id):
    lit = db.session.get(Literature, lit_id)
    if lit:
        db.session.delete(lit)
        db.session.commit()
        log_audit('DELETE', 'literature', lit_id, f'Deleted literature ID {lit_id}')
        flash('文献已删除。', 'success')
    else:
        flash('文献未找到。', 'warning')
    return redirect(url_for('literature_list'))


# =====================================================================
# Routes: Reaction Search (反应式检索)
# =====================================================================

@app.route('/reactions/search', methods=['GET', 'POST'])
@login_required
def reaction_search():
    form = ReactionSearchForm()
    results = []
    if form.validate_on_submit():
        query = Experiment.query.filter_by(is_active=True)
        if form.we_material.data:
            query = query.join(experiment_electrode).join(Electrode)\
                .filter(Electrode.electrode_type == 'WE',
                        Electrode.material.ilike(f'%{form.we_material.data}%'))
        if form.solvent.data:
            query = query.join(experiment_electrolyte).join(Electrolyte)\
                .filter(Electrolyte.solvent.ilike(f'%{form.solvent.data}%'))
        if form.electro_mode.data:
            query = query.filter(Experiment.electro_mode == form.electro_mode.data)
        if form.ph_min.data is not None:
            query = query.filter(Experiment.ph >= form.ph_min.data)
        if form.ph_max.data is not None:
            query = query.filter(Experiment.ph <= form.ph_max.data)
        if form.temp_min.data is not None:
            query = query.filter(Experiment.temperature >= form.temp_min.data)
        if form.temp_max.data is not None:
            query = query.filter(Experiment.temperature <= form.temp_max.data)
        if form.voltage_min.data is not None:
            query = query.filter(Experiment.potential >= form.voltage_min.data)
        if form.voltage_max.data is not None:
            query = query.filter(Experiment.potential <= form.voltage_max.data)
        if form.keyword.data:
            query = query.filter(db.or_(
                Experiment.name.ilike(f'%{form.keyword.data}%'),
                Experiment.notes.ilike(f'%{form.keyword.data}%'),
                Experiment.atmosphere.ilike(f'%{form.keyword.data}%')
            ))
        experiments = query.order_by(Experiment.created_at.desc()).all()
        if form.reactant_smiles.data or form.product_smiles.data:
            scored = _search_similar_reactions(
                form.reactant_smiles.data, form.product_smiles.data,
                experiments, form.sim_threshold.data or 0.3
            )
            results = scored
        else:
            results = [(exp, 1.0) for exp in experiments]
    return render_template('reactions/search.html', form=form, results=results)


@app.route('/api/reactions/similar', methods=['POST'])
@login_required
def api_reaction_similar():
    data = request.get_json()
    reactant = data.get('reactant_smiles', '')
    product = data.get('product_smiles', '')
    threshold = data.get('threshold', 0.3)
    limit = data.get('limit', 50)
    experiments = Experiment.query.filter_by(is_active=True).all()
    scored = _search_similar_reactions(reactant, product, experiments, threshold)
    scored = scored[:limit]
    return jsonify([{
        'id': exp.id, 'name': exp.name, 'similarity': round(score, 4),
        'reactant_smiles': exp.substrate_smiles,
        'product_smiles': exp.product_smiles,
        'doi': exp.doi, 'patent': exp.patent_number,
    } for exp, score in scored])


def _search_similar_reactions(reactant_smiles, product_smiles, experiments, threshold=0.3):
    """Search experiments by similarity of reactant/product SMILES."""
    reactant_fp = _smiles_fingerprint(reactant_smiles) if reactant_smiles else None
    product_fp = _smiles_fingerprint(product_smiles) if product_smiles else None
    results = []
    for exp in experiments:
        scores = []
        if reactant_fp and exp.substrate_smiles:
            fp = _smiles_fingerprint(exp.substrate_smiles)
            scores.append(_tanimoto_similarity(reactant_fp, fp))
        if product_fp and exp.product_smiles:
            fp = _smiles_fingerprint(exp.product_smiles)
            scores.append(_tanimoto_similarity(product_fp, fp))
        if not scores:
            scores.append(0.0)
        best = max(scores)
        if best >= threshold:
            results.append((exp, best))
    results.sort(key=lambda x: x[1], reverse=True)
    return results

# ------- Routes: Import / Export -------

@app.route('/import', methods=['GET', 'POST'])
@login_required
def import_data():
    form = ImportForm()
    results = None
    
    if form.validate_on_submit():
        file = form.file.data
        if not file:
            flash('请选择要导入的文件。', 'warning')
            return render_template('import.html', form=form)
        
        filename = secure_filename(file.filename)
        success_count = 0
        error_count = 0
        errors = []
        
        try:
            raw_bytes = file.read()
            
            if filename.endswith('.csv'):
                text = raw_bytes.decode('utf-8-sig')
                reader = csv.DictReader(StringIO(text))
                for i, row in enumerate(reader, 1):
                    try:
                        _import_experiment_from_dict(row, current_user.id)
                        success_count += 1
                    except Exception as e:
                        error_count += 1
                        errors.append(f'第{i}行: {str(e)}')
            
            elif filename.endswith('.json'):
                text = raw_bytes.decode('utf-8-sig')
                data = json.loads(text)
                if not isinstance(data, list):
                    data = [data]
                for i, record in enumerate(data, 1):
                    try:
                        _import_experiment_from_dict(record, current_user.id)
                        success_count += 1
                    except Exception as e:
                        error_count += 1
                        errors.append(f'第{i}条: {str(e)}')

            elif filename.endswith('.docx'):
                docx_results = _parse_docx_experiments(raw_bytes, current_user.id)
                success_count = docx_results['success']
                errors.extend(docx_results['errors'])
                error_count = docx_results['error_count']

            else:
                flash('仅支持 CSV、JSON 或 Word (.docx) 文件格式。', 'warning')
                return render_template('import.html', form=form)
            
            results = {'success': success_count, 'errors': errors, 'error_count': error_count}
            
            if success_count > 0:
                db.session.commit()
                log_audit('IMPORT', 'experiment', None, f'Imported {success_count} experiments from {filename}')
                flash(f'成功导入 {success_count} 条实验记录！', 'success')
            if error_count > 0:
                flash(f'{error_count} 条记录导入失败，请检查数据格式。', 'warning')
                
        except Exception as e:
            flash(f'导入过程出错: {str(e)}', 'danger')
    
    return render_template('import.html', form=form, results=results)


def _import_experiment_from_dict(data, user_id):
    """Import a single experiment record from a dict."""
    exp = Experiment(
        name=data.get('name') or data.get('实验名称', '未命名实验'),
        date_time=datetime.now(timezone.utc),
        user_id=user_id,
        batch_id=data.get('batch_id') or data.get('批次号'),
        cell_type=data.get('cell_type') or data.get('反应器类型'),
        cell_volume=_safe_float(data.get('cell_volume') or data.get('反应器体积')),
        separator=data.get('separator') or data.get('隔膜'),
        ph=_safe_float(data.get('ph') or data.get('pH')),
        temperature=_safe_float(data.get('temperature') or data.get('温度')),
        stirring_rpm=_safe_float(data.get('stirring_rpm') or data.get('搅拌速率')),
        atmosphere=data.get('atmosphere') or data.get('气氛'),
        feed_mode=data.get('feed_mode') or data.get('加料方式'),
        electro_mode=data.get('electro_mode') or data.get('电化学模式'),
        current=_safe_float(data.get('current') or data.get('电流')),
        potential=_safe_float(data.get('potential') or data.get('电位')),
        scan_rate=_safe_float(data.get('scan_rate') or data.get('扫描速率')),
        total_charge=_safe_float(data.get('total_charge') or data.get('总电荷量')),
        duration=_safe_float(data.get('duration') or data.get('实验时间')),
        operator_name=data.get('operator') or data.get('操作人员'),
        equipment_model=data.get('instrument_model') or data.get('仪器型号')
    )
    db.session.add(exp)
    db.session.flush()
    
    # Handle electrodes
    we = data.get('we_material') or data.get('工作电极材料')
    if we:
        eid = get_or_create_electrode(we, 'WE')
        if eid:
            db.session.execute(experiment_electrode.insert().values(
                experiment_id=exp.id, electrode_id=eid, function='WE'
            ))
    
    ce = data.get('ce_material') or data.get('对电极材料')
    if ce:
        eid = get_or_create_electrode(ce, 'CE')
        if eid:
            db.session.execute(experiment_electrode.insert().values(
                experiment_id=exp.id, electrode_id=eid, function='CE'
            ))
    
    # Handle electrolyte
    solvent = data.get('solvent') or data.get('溶剂')
    if solvent:
        salt = data.get('electrolyte') or data.get('支持电解质')
        conc = _safe_float(data.get('electrolyte_conc') or data.get('支持电解质浓度'))
        eid = get_or_create_electrolyte(solvent, salt, conc)
        if eid:
            db.session.execute(experiment_electrolyte.insert().values(
                experiment_id=exp.id, electrolyte_id=eid
            ))
    
    # Handle substrate
    sub = data.get('substrate') or data.get('底物')
    if sub:
        sub_conc = _safe_float(data.get('substrate_conc') or data.get('底物浓度'))
        sid = get_or_create_sample(sub, sub_conc, 'substrate')
        if sid:
            db.session.execute(experiment_sample.insert().values(
                experiment_id=exp.id, sample_id=sid, role='substrate'
            ))





def _parse_docx_experiments(raw_bytes, user_id):
    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(raw_bytes))
    records = []
    tables = doc.tables
    if tables:
        for table in tables:
            rows = table.rows
            if len(rows) < 2:
                continue
            header_cells = [c.text.strip() for c in rows[0].cells]
            has_header = _looks_like_header(header_cells)
            if has_header:
                for row_idx in range(1, len(rows)):
                    record = {}
                    cells = [c.text.strip() for c in rows[row_idx].cells]
                    for ci, cv in enumerate(cells):
                        if ci < len(header_cells) and cv:
                            record[_normalize_field_name(header_cells[ci])] = cv
                    if record:
                        records.append(record)
            else:
                record = {}
                for row in rows:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) >= 2 and cells[0]:
                        record[_normalize_field_name(cells[0].rstrip(':').rstrip('\uff1a'))] = cells[1]
                if record:
                    records.append(record)
    if not records:
        cur = {}
        for para in doc.paragraphs:
            t = para.text.strip()
            if not t:
                if cur:
                    records.append(cur)
                    cur = {}
                continue
            sep = '\uff1a' if '\uff1a' in t else (':' if ':' in t else None)
            if sep:
                parts = t.split(sep, 1)
                fn = _normalize_field_name(parts[0].strip())
                if fn:
                    cur[fn] = parts[1].strip()
            else:
                cur['notes'] = cur.get('notes', '') + '\n' + t
        if cur:
            records.append(cur)
    success_count = 0
    error_count = 0
    errors = []
    for i, rec in enumerate(records, 1):
        try:
            _import_experiment_from_dict(rec, user_id)
            success_count += 1
        except Exception as e:
            error_count += 1
            errors.append(f'\u7b2c{i}\u6761: {str(e)}')
    return {'success': success_count, 'errors': errors, 'error_count': error_count}


_FIELD_MAP = {
    '\u5b9e\u9a8c\u540d\u79f0': 'name', 'name': 'name',
    '\u5de5\u4f5c\u7535\u6781\u6750\u6599': 'we_material', 'we_material': 'we_material',
    '\u5de5\u4f5c\u7535\u6781': 'we_material',
    '\u5de5\u4f5c\u7535\u6781\u9762\u79ef': 'we_area', 'we_area': 'we_area',
    'we_pretreatment': 'we_pretreatment',
    '\u53c2\u6bd4\u7535\u6781': 're_type', 're_type': 're_type',
    '\u5bf9\u7535\u6781\u6750\u6599': 'ce_material', 'ce_material': 'ce_material',
    '\u5bf9\u7535\u6781': 'ce_material',
    '\u6eb6\u5242': 'solvent', 'solvent': 'solvent',
    '\u652f\u6301\u7535\u89e3\u8d28': 'electrolyte', 'electrolyte': 'electrolyte',
    'electrolyte_conc': 'electrolyte_conc',
    '\u5e95\u7269': 'substrate', 'substrate': 'substrate',
    'substrate_conc': 'substrate_conc',
    'ph': 'ph', 'pH': 'ph', 'PH': 'ph',
    '\u6e29\u5ea6': 'temperature', 'temperature': 'temperature',
    '\u6405\u62cc\u901f\u7387': 'stirring_rpm', 'stirring_rpm': 'stirring_rpm',
    '\u6c14\u6c1b': 'atmosphere', 'atmosphere': 'atmosphere',
    '\u53cd\u5e94\u5668\u7c7b\u578b': 'cell_type', 'cell_type': 'cell_type',
    'cell_volume': 'cell_volume',
    '\u9694\u819c': 'separator', 'separator': 'separator',
    'feed_mode': 'feed_mode',
    '\u7535\u5316\u5b66\u6a21\u5f0f': 'electro_mode', 'electro_mode': 'electro_mode',
    '\u7535\u6d41': 'current', 'current': 'current',
    '\u7535\u4f4d': 'potential', 'potential': 'potential', '\u7535\u538b': 'potential',
    'scan_rate': 'scan_rate',
    '\u603b\u7535\u8377\u91cf': 'total_charge', 'total_charge': 'total_charge',
    '\u5b9e\u9a8c\u65f6\u95f4': 'duration', 'duration': 'duration',
    '\u50ac\u5316\u5242': 'catalyst', 'catalyst': 'catalyst',
    'catalyst_load': 'catalyst_load',
    '\u6dfb\u52a0\u5242': 'additives', 'additives': 'additives',
    '\u5149\u6e90': 'light_source', 'light_source': 'light_source',
    'light_intensity': 'light_intensity',
    'measurement_method': 'measurement_method',
    'sampling_rate': 'sampling_rate',
    'analysis_method': 'analysis_method',
    '\u4eea\u5668\u578b\u53f7': 'instrument_model', 'instrument_model': 'instrument_model',
    '\u8bbe\u5907\u578b\u53f7': 'instrument_model',
    '\u64cd\u4f5c\u4eba\u5458': 'operator', 'operator': 'operator',
    '\u5b9e\u9a8c\u8005': 'operator',
    'batch_id': 'batch_id',
    'safety_info': 'safety_info',
    'waste_category': 'waste_category',
    '\u5907\u6ce8': 'notes', 'notes': 'notes',
}


def _normalize_field_name(name):
    name = name.strip().rstrip(':').rstrip('\uff1a')
    if name in _FIELD_MAP:
        return _FIELD_MAP[name]
    for k, v in _FIELD_MAP.items():
        if isinstance(k, str) and k and (k in name or name in k):
            return v
    return name


def _looks_like_header(cells):
    indicators = ['name', 'experiment', 'we_', 'solvent', 'substrate', 'temperature', 'ph']
    for cell in cells:
        cl = cell.lower().strip()
        for ind in indicators:
            if ind in cl:
                return True
    if cells and sum(len(c) for c in cells) / len(cells) < 20:
        return True
    return False
def _safe_float(val):
    """Safely convert a value to float, returning None if not possible."""
    if val is None:
        return None
    try:
        return float(val)
    except (ValueError, TypeError):
        return None


@app.route('/export')
@login_required
def export_data():
    fmt = request.args.get('format', 'csv')
    exp_ids = request.args.get('ids', '')
    
    query = Experiment.query.filter_by(is_active=True)
    if exp_ids:
        ids = [int(x) for x in exp_ids.split(',') if x.strip().isdigit()]
        if ids:
            query = query.filter(Experiment.id.in_(ids))
    
    experiments = query.order_by(Experiment.created_at.desc()).all()
    
    if fmt == 'csv':
        output = StringIO()
        writer = csv.writer(output)
        headers = [
            '实验编号', '实验名称', '日期时间', '工作电极材料', '工作电极面积',
            '参比电极', '对电极材料', '溶剂', '支持电解质', '支持电解质浓度',
            '底物', '底物浓度', 'pH', '温度', '搅拌速率', '气氛',
            '反应器类型', '隔膜', '电化学模式', '电流', '电位', '扫描速率',
            '总电荷量', '实验时间', '催化剂', '操作人员', '仪器型号', '备注'
        ]
        writer.writerow(headers)
        
        for exp in experiments:
            we_mat = ''
            ce_mat = ''
            re_type = ''
            for elec in exp.electrodes:
                funcs = db.session.execute(
                    db.select(experiment_electrode.c.function).where(
                        experiment_electrode.c.experiment_id == exp.id,
                        experiment_electrode.c.electrode_id == elec.id
                    )
                ).scalars().all()
                for f in funcs:
                    if f == 'WE': we_mat = elec.material
                    elif f == 'CE': ce_mat = elec.material
                    elif f == 'RE': re_type = elec.material
            
            solvent = ''
            electrolyte = ''
            electrolyte_conc = ''
            for elec in exp.electrolytes:
                solvent = elec.solvent
                electrolyte = elec.supporting_salt
                electrolyte_conc = elec.concentration
            
            substrate = ''
            substrate_conc = ''
            catalyst = ''
            for samp in exp.samples:
                roles = db.session.execute(
                    db.select(experiment_sample.c.role).where(
                        experiment_sample.c.experiment_id == exp.id,
                        experiment_sample.c.sample_id == samp.id
                    )
                ).scalars().all()
                for r in roles:
                    if r == 'substrate':
                        substrate = samp.name
                        substrate_conc = samp.concentration
                    elif r == 'catalyst':
                        catalyst = samp.name
            
            writer.writerow([
                exp.id, exp.name, exp.date_time,
                we_mat, '', re_type, ce_mat,
                solvent, electrolyte, electrolyte_conc,
                substrate, substrate_conc,
                exp.ph, exp.temperature, exp.stirring_rpm, exp.atmosphere,
                exp.cell_type, exp.separator,
                exp.electro_mode, exp.current, exp.potential, exp.scan_rate,
                exp.total_charge, f'{exp.duration}{exp.duration_unit or ""}',
                catalyst, exp.operator_name, exp.equipment_model, exp.notes
            ])
        
        output.seek(0)
        response = make_response(output.getvalue())
        response.headers['Content-Disposition'] = 'attachment; filename=experiments_export.csv'
        response.headers['Content-Type'] = 'text/csv; charset=utf-8-sig'
        return response
    
    elif fmt == 'json':
        data = []
        for exp in experiments:
            elec_list = [{'material': e.material, 'type': e.electrode_type} for e in exp.electrodes]
            elec_list_elec = [{'solvent': e.solvent, 'salt': e.supporting_salt} for e in exp.electrolytes]
            samp_list = [{'name': s.name} for s in exp.samples]
            
            data.append({
                'id': exp.id,
                'name': exp.name,
                'date_time': exp.date_time.isoformat() if exp.date_time else None,
                'electrodes': elec_list,
                'electrolytes': elec_list_elec,
                'samples': samp_list,
                'ph': exp.ph,
                'temperature': exp.temperature,
                'electro_mode': exp.electro_mode,
                'current': exp.current,
                'potential': exp.potential,
                'duration': exp.duration,
                'notes': exp.notes
            })
        
        response = make_response(json.dumps(data, ensure_ascii=False, indent=2))
        response.headers['Content-Disposition'] = 'attachment; filename=experiments_export.json'
        response.headers['Content-Type'] = 'application/json; charset=utf-8'
        return response
    
    flash('不支持的导出格式。', 'warning')
    return redirect(url_for('experiment_list'))


# Download template CSV
@app.route('/export/template')
@login_required
def export_template():
    """Download a Word (.docx) import template with field guide and example data."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    from io import BytesIO

    def _shd(cell, ch):
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ch}"/>'))

    def _wc(cell, text, bold=False, sz=10, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(sz)
        run.font.name = "SimSun"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # Title
    t = doc.add_heading("", level=0).add_run("电化学实验数据导入模板")
    t.font.size = Pt(22)
    t.font.color.rgb = RGBColor(26, 35, 126)
    t.font.name = "SimHei"
    t._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Electrochemical Experimental Data Import Template")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(100, 100, 100)
    r.font.italic = True

    doc.add_paragraph(chr(8212) * 40)

    # Instructions
    ih = doc.add_heading("", level=1).add_run("使用说明")
    ih.font.color.rgb = RGBColor(26, 35, 126)
    ih.font.name = "SimHei"
    ih._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    instrs = [
        "1. 本模板包含电化学实验的所有标准字段，按类别分组排列。",
        "2. 已包含三个典型实验的示例数据（水电解制氢、CO₂还原、有机电合成）。",
        "3. 填写您的实验数据时，将示例行替换为您的实验数据即可。",
        "4. 请保留第一行（字段名行）不变，系统导入时依据该行识别数据内容。",
        "5. 字段名支持中文或英文，系统会自动匹配。未填写的字段留空即可。",
        "6. 必填字段：实验名称、工作电极材料、溶剂、底物、电化学模式。",
        "7. 完成后将文件保存为 .docx 格式，然后在平台导入页面上传。",
    ]
    for t in instrs:
        pp = doc.add_paragraph(t)
        pp.paragraph_format.space_before = Pt(2)
        pp.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # Table
    th = doc.add_heading("", level=1).add_run("实验数据表")
    th.font.color.rgb = RGBColor(26, 35, 126)
    th.font.name = "SimHei"
    th._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    cats = [
        ("基本信息", [
            ("实验名称", "name", "文本"),
            ("时间日期", "date_time", "日期时间"),
            ("批次号", "batch_id", "文本"),
            ("操作人员", "operator", "文本"),
        ]),
        ("电极参数", [
            ("工作电极材料", "we_material", "文本"),
            ("工作电极面积", "we_area", "cm²"),
            ("工作电极形貌", "we_geometry", "文本"),
            ("工作电极预处理", "we_pretreatment", "文本"),
            ("参比电极类型", "re_type", "文本"),
            ("对电极材料", "ce_material", "文本"),
        ]),
        ("溶液体系参数", [
            ("溶剂", "solvent", "文本"),
            ("支持电解质", "electrolyte", "文本"),
            ("支持电解质浓度", "electrolyte_conc", "M"),
            ("底物", "substrate", "文本"),
            ("底物浓度", "substrate_conc", "M"),
            ("pH", "pH", "0-14"),
            ("温度", "temperature", "℃"),
            ("搅拌速率", "stirring_rpm", "rpm"),
            ("气氛", "atmosphere", "文本"),
            ("加料方式", "feed_mode", "文本"),
            ("加料速率", "feed_rate", "mL/min"),
        ]),
        ("反应器与环境", [
            ("反应器类型", "cell_type", "文本"),
            ("反应器体积", "cell_volume", "mL"),
            ("隔膜", "separator", "文本"),
        ]),
        ("电化学参数", [
            ("电化学模式", "electro_mode", "枚举"),
            ("电流", "current", "A"),
            ("电位", "potential", "V"),
            ("扫描速率", "scan_rate", "mV/s"),
            ("总电荷量", "total_charge", "C"),
            ("实验时间", "duration", ""),
            ("时间单位", "duration_unit", "h/s"),
        ]),
        ("催化剂与光照", [
            ("催化剂", "catalyst", "文本"),
            ("催化剂负载", "catalyst_load", "mg"),
            ("添加剂", "additives", "文本"),
            ("光源类型", "light_source", "文本"),
            ("光照强度", "light_intensity", "mW/cm²"),
            ("光照时间", "light_duration", "s"),
        ]),
        ("测量与分析", [
            ("测量方法", "measurement_method", "文本"),
            ("采样频率", "sampling_rate", "Hz"),
            ("产物分析方法", "analysis_method", "文本"),
            ("不确定度", "uncertainty", "%"),
        ]),
        ("其他信息", [
            ("仪器型号", "instrument_model", "文本"),
            ("安全信息", "safety_info", "文本"),
            ("废弃物分类", "waste_category", "文本"),
            ("备注", "notes", "文本"),
        ]),
    ]

    ex1 = {"实验名称":"HER产氢测试","时间日期":"2026-07-01 10:00","批次号":"BATCH-2026-001","操作人员":"张三","工作电极材料":"Ni泡沫","工作电极面积":"1.0","工作电极形貌":"泡沫状","工作电极预处理":"机械抛光+超声清洗","参比电极类型":"Ag/AgCl (KCl饱和)","对电极材料":"Ni泡沫","溶剂":"1.0 M KOH水溶液","支持电解质":"KOH","支持电解质浓度":"1.0","底物":"纯水 (H₂O)","pH":"13","温度":"25","搅拌速率":"600","气氛":"氙气惰性保护","反应器类型":"H型电解池","反应器体积":"100","隔膜":"Nafion 117","电化学模式":"恒电流","电流":"0.1","总电荷量":"360","实验时间":"6","时间单位":"h","催化剂":"NiMo合金","催化剂负载":"5","测量方法":"CP","产物分析方法":"气相色谱(GC)","仪器型号":"CHI660E","备注":"碱性HER性能测试"}

    ex2 = {"实验名称":"CO₂还原产乙烯","时间日期":"2026-07-02 14:30","批次号":"BATCH-2026-002","操作人员":"李四","工作电极材料":"铜箔","工作电极面积":"0.5","工作电极形貌":"平板","工作电极预处理":"化学刻刻(CuSO₄溶液)","参比电极类型":"Ag/AgCl (KCl饱和)","对电极材料":"铂片","溶剂":"0.1 M KHCO₃水溶液","支持电解质":"KHCO₃","支持电解质浓度":"0.1","底物":"CO₂ (饱和)","pH":"6.8","温度":"25","搅拌速率":"600","气氛":"CO₂持续通入 (20 mL/min)","反应器类型":"H型电解池","反应器体积":"100","隔膜":"Nafion 117","电化学模式":"线性扫描","电位":"-1.0","扫描速率":"50","催化剂":"Cu纳米颗粒","催化剂负载":"1","测量方法":"LSV","产物分析方法":"气相色谱(GC)","仪器型号":"CHI660E","备注":"CO₂电还原制乙烯"}

    ex3 = {"实验名称":"苯磺酰胺氧化合成","时间日期":"2026-07-03 09:00","批次号":"BATCH-2026-003","操作人员":"王五","工作电极材料":"石墨烯沉积Pt","工作电极面积":"1.0","工作电极形貌":"平板","工作电极预处理":"热处理+表面修饰","参比电极类型":"Ag/AgCl (KCl饱和)","对电极材料":"铂片","溶剂":"MeCN (乙腈)","支持电解质":"TBAPF₆","支持电解质浓度":"0.1","底物":"苯磺酰胺","底物浓度":"0.1","pH":"7.0 (磷酸盐缓冲)","温度":"25","搅拌速率":"0","气氛":"氙气惰性保护","反应器类型":"单室玻璃瓶","反应器体积":"50","隔膜":"无隔膜","电化学模式":"恒电位","电位":"-2.0","总电荷量":"200","实验时间":"1.5","时间单位":"h","添加剂":"K₂CO₃ 0.01 M","光源类型":"LED 365 nm","光照强度":"50","光照时间":"5400","测量方法":"CP","产物分析方法":"HPLC","仪器型号":"CHI660E","备注":"光电催化有机合成"}

    NUM_COLS = 7
    NUM_ROWS = 1 + sum(len(f[1]) for f in cats) + len(cats)
    table = doc.add_table(rows=NUM_ROWS, cols=NUM_COLS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    widths_cm = [Cm(2.0), Cm(3.2), Cm(3.5), Cm(2.5), Cm(3.5), Cm(3.5), Cm(3.5)]
    widths_twips = [int(w.emu / 914400 * 1440) for w in widths_cm]
    tbl_pr = table._tbl.tblPr
    tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="9072" w:type="dxa"/>')
    tbl_pr.append(tbl_w)

    hdrs = ["类别", "字段名称", "字段标识", "类型/单位", "示例1: 水电解制氢", "示例2: CO₂还原", "示例3: 苯磺酰胺氧化"]
    hcols = ["1a237e", "1a237e", "1a237e", "1a237e", "283593", "283593", "283593"]

    for ci, (hdr, hc) in enumerate(zip(hdrs, hcols)):
        cell = table.rows[0].cells[ci]
        _wc(cell, hdr, bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))
        _shd(cell, hc)

    ri = 0
    for cat_name, fields in cats:
        ri += 1
        row = table.rows[ri]
        _wc(row.cells[0], cat_name, bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))
        _shd(row.cells[0], "3949ab")
        for ci in range(1, NUM_COLS):
            _wc(row.cells[ci], "", sz=9)
            _shd(row.cells[ci], "e8eaf6")

        for fname, fid, ftype in fields:
            ri += 1
            row = table.rows[ri]
            v1 = ex1.get(fname, "")
            v2 = ex2.get(fname, "")
            v3 = ex3.get(fname, "")
            _wc(row.cells[0], "", sz=9)
            _wc(row.cells[1], fname, bold=True, sz=9)
            _wc(row.cells[2], fid, sz=8, color=(100, 100, 100))
            _wc(row.cells[3], ftype, sz=8, color=(100, 100, 100))
            _wc(row.cells[4], v1, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _wc(row.cells[5], v2, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _wc(row.cells[6], v3, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            if ri % 2 == 0:
                for ci in range(NUM_COLS):
                    _shd(row.cells[ci], "f5f5f5")

    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("— 电化学实验条件数据库平台 (EChemDB) —")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(150, 150, 150)

    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="ecdb_template.docx")


# ------- Routes: API (JSON) -------

@app.route('/api/experiments')
@login_required
def api_experiments():
    experiments = Experiment.query.filter_by(is_active=True)\
        .order_by(Experiment.created_at.desc()).limit(100).all()
    data = []
    for exp in experiments:
        data.append({
            'id': exp.id,
            'name': exp.name,
            'date_time': exp.date_time.isoformat() if exp.date_time else None,
            'electro_mode': exp.electro_mode,
            'temperature': exp.temperature,
            'ph': exp.ph,
            'current': exp.current,
            'potential': exp.potential,
            'we_material': [e.material for e in exp.electrodes if e.electrode_type == 'WE'],
            'solvent': [e.solvent for e in exp.electrolytes],
        })
    return jsonify(data)


@app.route('/api/stats/summary')
@login_required
def api_stats_summary():
    total = Experiment.query.filter_by(is_active=True).count()
    with_analysis = Analysis.query.count()
    
    avg_yield = db.session.query(func.avg(Analysis.yield_pct)).scalar() or 0
    avg_fe = db.session.query(func.avg(Analysis.faraday_efficiency)).scalar() or 0
    
    return jsonify({
        'total_experiments': total,
        'with_analysis': with_analysis,
        'average_yield': round(float(avg_yield), 2),
        'average_faraday_efficiency': round(float(avg_fe), 2)
    })




_SMILES_NAMES = {
    "O": "Water",
    "C=O": "Formaldehyde",
    "CO": "Methanol",
    "CCO": "Ethanol",
    "O=C=O": "Carbon Dioxide",
    "CC=O": "Acetaldehyde",
    "c1ccccc1": "Benzene",
    "CC": "Ethane",
    "C=C": "Ethylene",
    "CC=C": "Propylene",
    "N": "Ammonia",
    "Cl": "HCl",
    "C#N": "HCN",
    "CC(O)=O": "Acetic Acid",
    "CCOC(C)=O": "Ethyl Acetate",
    "CC(C)=O": "Acetone",
}

@app.route('/api/smiles/name', methods=['POST'])
@login_required
def api_smiles_name():
    data = request.get_json()
    smi = data.get('smiles', '').strip()
    if not smi:
        return jsonify({'name': '', 'smiles': ''})
    name = _SMILES_NAMES.get(smi)
    if name:
        return jsonify({'name': name, 'smiles': smi, 'source': 'lookup'})
    try:
        import requests as req
        from urllib.parse import quote
        url = 'https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/smiles/' + quote(smi) + '/property/IUPACName,TITLE/JSON'
        resp = req.get(url, timeout=5)
        if resp.status_code == 200:
            props = resp.json().get('PropertyTable', {}).get('Properties', [{}])[0]
            name = props.get('IUPACName') or props.get('Title', '')
            if name:
                return jsonify({'name': name, 'smiles': smi, 'source': 'pubchem'})
    except Exception:
        pass
    return jsonify({'name': 'Molecule: ' + smi[:50], 'smiles': smi, 'source': 'smiles'})



def _smiles_fingerprint(smiles):
    """Generate a simple structural fingerprint from SMILES.
    Counts atom types, bond types, ring features, and functional groups."""
    if not smiles:
        return {}
    fp = {}
    elements = {'C','N','O','S','P','F','Cl','Br','I','H'}
    for elem in elements:
        fp['elem_' + elem] = smiles.count(elem)
    bond_types = {'=', '#', ':', '.'}
    for b in bond_types:
        fp['bond_' + b] = smiles.count(b)
    fp['rings'] = smiles.count('1') + smiles.count('2') + smiles.count('3')
    fp['branches'] = smiles.count('(') + smiles.count(')')
    fp['charges'] = smiles.count('+') + smiles.count('-')
    fp['aromatic'] = smiles.count('c') + smiles.count('C') - smiles.count('C(') + smiles.count('n')
    fp['length'] = len(smiles)
    total = sum(abs(v) for v in fp.values()) or 1
    return {k: v/total for k, v in fp.items()}


def _tanimoto_similarity(fp1, fp2):
    """Compute Tanimoto coefficient between two fingerprints."""
    all_keys = set(fp1.keys()) | set(fp2.keys())
    dot = sum(fp1.get(k, 0) * fp2.get(k, 0) for k in all_keys)
    norm1 = sum(v*v for v in fp1.values())
    norm2 = sum(v*v for v in fp2.values())
    if norm1 + norm2 - dot == 0:
        return 0.0
    return dot / (norm1 + norm2 - dot)


def _search_similar_structures(query_smiles, experiments, threshold=0.3):
    """Search experiments with SMILES similar to the query."""
    query_fp = _smiles_fingerprint(query_smiles)
    if not query_fp:
        return []
    results = []
    for exp in experiments:
        best_score = 0.0
        candidates = []
        if exp.substrate_smiles:
            candidates.append(exp.substrate_smiles)
        if exp.product_smiles:
            candidates.append(exp.product_smiles)
        if exp.smiles:
            candidates.append(exp.smiles)
        for smi in candidates:
            fp = _smiles_fingerprint(smi)
            score = _tanimoto_similarity(query_fp, fp)
            if score > best_score:
                best_score = score
        if best_score >= threshold:
            results.append((exp, best_score))
    results.sort(key=lambda x: x[1], reverse=True)
    return results


_SMILES_NAMES = {
    "O": "Water", "C=O": "Formaldehyde", "CO": "Methanol", "CCO": "Ethanol",
    "O=C=O": "Carbon Dioxide", "CC=O": "Acetaldehyde", "c1ccccc1": "Benzene",
    "CC": "Ethane", "C=C": "Ethylene", "CC=C": "Propylene", "N": "Ammonia",
    "Cl": "HCl", "Br": "HBr", "CC(O)=O": "Acetic Acid",
    "CCOC(C)=O": "Ethyl Acetate", "CC(C)=O": "Acetone",
}

@app.route('/struct-img/<path:smiles>')
def struct_img(smiles):
    import requests as req
    from flask import Response
    try:
        url = 'https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/smiles/' + smiles + '/PNG'
        resp = req.get(url, timeout=10)
        if resp.status_code == 200:
            return Response(resp.content, mimetype='image/png')
    except Exception:
        pass
    svg = '<svg xmlns="http://www.w3.org/2000/svg" width="180" height="100"><text x="10" y="55" font-family="monospace" font-size="11" fill="#888">' + smiles[:50] + '</text></svg>'
    return Response(svg, mimetype='image/svg+xml')
@app.route('/health')
def health():
    """Health check endpoint for monitoring."""
    return jsonify({
        'status': 'ok',
        'experiments': Experiment.query.filter_by(is_active=True).count(),
        'literature': Literature.query.count(),
    })


# ------- Error Handlers -------

@app.errorhandler(404)
def not_found(e):
    return render_template('base.html', error='404 - 页面未找到'), 404


@app.errorhandler(500)
def server_error(e):
    return render_template('base.html', error='500 - 服务器内部错误'), 500


# ------- Main -------

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    import argparse
    import os
    parser = argparse.ArgumentParser(description="EChemDB - Electrochemical Database Platform")
    parser.add_argument("--port", type=int, default=int(os.environ.get("PORT", 5002)), help="Server port")
    parser.add_argument("--host", type=str, default="0.0.0.0", help="Bind address")
    parser.add_argument("--debug", action="store_true", default=False, help="Enable debug mode")
    args = parser.parse_args()
    app.run(host=args.host, port=args.port, debug=args.debug)
